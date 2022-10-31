# -*- coding:utf-8-*-
# Author: Xin Ge
# Created on 19.09.2022

from docx import Document
import os
from pprint import pp


'''
    pass parameters to function
'''    
def fillDoc2(curr_dir, info_list):

    pp(curr_dir)
    for i in info_list:
        try:
            doc_name = i[0] + i[1] + '.docx'
            input_dir = os.path.join(curr_dir, 'input', doc_name)
            output_dir = os.path.join(curr_dir, 'output', doc_name )
            document = Document(input_dir) #读入文件
            tables = document.tables #获取文件中的表格集
            table = tables[0]#获取文件中的第0个表格
            table.cell(4,1).text = i[0]
            
            for j in document.paragraphs:
                ptext = j.text.replace(' ','')
                print (ptext)
                
                #if '鉴定编号' in ptext:
                #    j.text = j.text.rstrip() + i[1]
                
                if '报告编号' in ptext:
                    j.text = j.text + i[1]
                    print('替换》' + j.text)

            document.save(output_dir)
        except Exception as e:
            print(e)
                
            
if __name__ == '__main__':
    fillDoc2()
